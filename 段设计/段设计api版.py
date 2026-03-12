import os
import json
import re
import math
import numpy as np
import requests
from CoolProp.CoolProp import PropsSI
from docx import Document
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
import traceback
import time
import streamlit as st
import pandas as pd
import subprocess

# ===================== 1. 核心配置（仅删除AI_CONFIG，其余保留） =====================
# 存储目录（适配Streamlit临时目录）
STORAGE = {
    "design_results": "/tmp/design_results",
    "reports": "/tmp/reports"
}
for dir_path in STORAGE.values():
    os.makedirs(dir_path, exist_ok=True)

# 阶段映射
STAGE_MAP = {
    0: "参数配置中",
    1: "处理数据中",
    2: "优化计算中",
    3: "生成报告中",
    4: "完成",
    5: "失败"
}

# 预定义工质（供选择框用）
WORKING_MEDIA_OPTIONS = {
    "空气": "Air",
    "氮气": "Nitrogen",
    "氦气": "Helium",
    "二氧化碳": "CarbonDioxide",
    "氧气": "Oxygen"
}

# 预定义模板（完全保留原结构）
SECTION_TEMPLATES = {
    "default": {
        "compressed_coef": 0.875,
        "cooling_coef": 0.98,
        "efficient_pol": 0.80,
        "pressure_ratio": 2.4,
        "angle_outlet": 50,
        "discharge_coef": 0.26,
        "blade_number": 21,
        "loss": 0.025,
        "U2": 300.0,
        "flange_length": 300,
        "flange_width": 320,
        "diameter_axis": 0.13,
        "Kd": 1.03,
        "add_Qv": 0.1,
        "add_T": 0.0,
        "stages": []
    },
    "high_efficiency": {
        "compressed_coef": 0.90,
        "cooling_coef": 0.98,
        "efficient_pol": 0.85,
        "pressure_ratio": 2.0,
        "angle_outlet": 45,
        "discharge_coef": 0.28,
        "blade_number": 24,
        "loss": 0.02,
        "U2": 320.0,
        "flange_length": 310,
        "flange_width": 330,
        "diameter_axis": 0.125,
        "Kd": 1.05,
        "add_Qv": 0.08,
        "add_T": 0.0,
        "stages": []
    }
}
STAGE_TEMPLATES = {
    "default": {
        "loss": 0.025,
        "efficient_pol": 0.80,
        "thick_blade": 4,
        "blockage_coef_inlet": 0.85,
        "blockage_coef_outlet": 0.95,
        "width_diameter": 0.04,
        "Kc": 1.05,
        "Kv0": 1.2,
        "Ds": 0.35
    },
    "high_pressure": {
        "loss": 0.03,
        "efficient_pol": 0.78,
        "thick_blade": 5,
        "blockage_coef_inlet": 0.82,
        "blockage_coef_outlet": 0.93,
        "width_diameter": 0.035,
        "Kc": 1.08,
        "Kv0": 1.5,
        "Ds": 0.32
    }
}


# ===================== 2. 核心类与函数（完全保留原逻辑，无任何修改） =====================
# 1. 参数生成类
class ParameterGenerator:
    def __init__(self):
        self.parameters = {
            'boundary_condition': {'volume_flow_rate_inlet': None, 'standard_volume_rate': 1, 'temperature_inlet': None,
                                   'pressure_inlet': None, 'pressure_outlet': None},
            'working_medium': {'k_index': None, 'R_index': None, 'rotation_speed': None, 'seal_teeth': 5,
                               'seal_gap': 0.4, 'component': []},
            'sections': []
        }

    def generate_section_stage(self, section_type="default", stage_type="default", num_stages=3, section_index=None):
        section = SECTION_TEMPLATES.get(section_type, SECTION_TEMPLATES["default"]).copy()
        stages = []
        for i in range(num_stages):
            stage = STAGE_TEMPLATES.get(stage_type, STAGE_TEMPLATES["default"]).copy()
            stage["stage_id"] = f"stage_{len(self.parameters['sections'])}_{i}"
            stages.append(stage)
        section["stages"] = stages
        section["section_name"] = f"Section_{len(self.parameters['sections']) + 1}"
        section["total_stage"] = num_stages
        if section_index is None:
            self.parameters['sections'].append(section)
        else:
            self.parameters['sections'].insert(section_index, section)
        return f"成功添加 {section_type} 类型的段，包含 {num_stages} 个 {stage_type} 类型的级"

    def parse_boundary_conditions(self, user_input):
        extracted = {}
        conditions = {
            'volume_flow_rate_inlet': r'(入口流量|体积流量|流量)[=:]*\s*(\d+\.?\d*)\S*(Nm3/h|m³/h)',
            'temperature_inlet': r'(入口温度|进口温度|温度)[=:]*\s*(\d+\.?\d*)\S*K',
            'pressure_inlet': r'(入口压力|进口压力|压力)[=:]*\s*(\d+\.?\d*)\S*MPa',
            'pressure_outlet': r'(出口压力|出口压力)[=:]*\s*(\d+\.?\d*)\S*MPa',
            'rotation_speed': r'(转速|转数|旋转速度)[=:]*\s*(\d+\.?\d*)\S*rpm',
            'standard_volume_rate': r'(标准体积流量|标准流量|标准状态下的体积流量)[=:]*\s*(\d+\.?\d*)\S*(Nm3/h|m3/h)'
        }
        missing_conditions = []
        for key, pattern in conditions.items():
            match = re.search(pattern, user_input, re.IGNORECASE)
            if match:
                extracted[key] = float(match.group(2))
            else:
                if key not in ['volume_flow_rate_inlet', 'rotation_speed']:
                    missing_conditions.append(key)
        required_conditions = ['standard_volume_rate', 'temperature_inlet', 'pressure_inlet', 'pressure_outlet']
        missing_required = [c for c in required_conditions if c not in extracted]
        if missing_required:
            return (False, f"缺少必要的边界条件: {', '.join(missing_required)}")
        for key, value in extracted.items():
            if key in self.parameters['boundary_condition']:
                self.parameters['boundary_condition'][key] = value
            elif key in self.parameters['working_medium']:
                self.parameters['working_medium'][key] = value
        if 'rotation_speed' not in extracted:
            self.parameters['working_medium']['rotation_speed'] = 9500.0
        return (True, "边界条件解析成功")

    def calculate_working_medium_properties(self, user_input):
        coolprop_names = {
            '氦': 'Helium', '氦气': 'Helium', '氮': 'Nitrogen', '氮气': 'Nitrogen',
            'N2': 'Nitrogen', '氧': 'Oxygen', '氧气': 'Oxygen', 'O2': 'Oxygen',
            '氢': 'Hydrogen', '氢气': 'Hydrogen', '氩': 'Argon', '氩气': 'Argon',
            '二氧化碳': 'CarbonDioxide', 'CO2': 'CarbonDioxide', '甲烷': 'Methane',
            'CH4': 'Methane', '乙烷': 'Ethane', 'C2H6': 'Ethane', '丙烷': 'Propane',
            'C3H8': 'Propane', '丁烷': 'Butane', 'C4H10': 'Butane', '空气': 'Air',
            '水': 'Water', '水蒸气': 'Water'
        }
        parts = user_input.split(';')
        components = []
        T_in = None
        P_in = None
        for part in parts:
            if '进口温度' in part:
                match = re.search(r'=(\d+\.?\d*)', part)
                if match:
                    T_in = float(match.group(1))
            elif '进口压力' in part:
                match = re.search(r'=(\d+\.?\d*)', part)
                if match:
                    P_in = float(match.group(1)) * 1e6
            elif 'CoolProp代名为' in part:
                match_name = re.search(r'CoolProp代名为=(\w+)', part)
                match_index = re.search(r'组分(\d+)', part)
                if match_name and match_index:
                    components.append({'index': match_index.group(1), 'name': match_name.group(1)})
            elif '摩尔分数' in part and '组分' in part:
                match_index = re.search(r'组分(\d+)', part)
                match_value = re.search(r'=(\d+\.?\d*)', part)
                if match_index and match_value:
                    comp_index = match_index.group(1)
                    mole_fraction = float(match_value.group(1)) / 100.0
                    for comp in components:
                        if comp['index'] == comp_index:
                            comp['mole_fraction'] = mole_fraction
                            break
        if T_in is None or P_in is None:
            return (False, "未能提取进口温度或压力")
        if not components:
            return (False, "未能提取组分信息")
        total_mole = sum(comp.get('mole_fraction', 0) for comp in components)
        if abs(total_mole - 1.0) > 0.01:
            return (False, f"摩尔分数总和为{total_mole}，不等于1")
        try:
            M_avg = 0.0
            for comp in components:
                M_i = PropsSI('M', '', 0, '', 0, comp['name'])
                comp['M'] = M_i
                M_avg += comp['mole_fraction'] * M_i
            R_specific = 8314.462618 / M_avg / 1000
            mix_string = "&".join([f"{comp['name']}[{comp['mole_fraction']}]" for comp in components])
            Cp_mass = PropsSI('CP0MASS', 'T', T_in, 'P', P_in, mix_string)
            Cv_mass = Cp_mass - R_specific
            k_index = Cp_mass / Cv_mass
            rho_index = PropsSI('D', 'T', T_in, 'P', P_in, mix_string)
            self.parameters['working_medium']['k_index'] = k_index
            self.parameters['working_medium']['R_index'] = R_specific
            self.parameters['working_medium']['rho_index'] = rho_index
            self.parameters['working_medium']['component'] = [
                (comp['M'], comp['mole_fraction'], PropsSI('CPMASS', 'T', T_in, 'P', P_in, comp['name']))
                for comp in components
            ]
            return (True, f"物性计算成功: k={k_index:.4f}, R={R_specific:.2f}")
        except Exception as e:
            return (False, f"物性计算失败: {str(e)}")

    def save_parameters(self, file_path):
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(self.parameters, f, indent=4, ensure_ascii=False)
            return (True, f"参数已保存到 {file_path}")
        except Exception as e:
            return (False, f"保存失败: {str(e)}")

    def validate_parameters(self):
        missing = []
        bc = self.parameters['boundary_condition']
        for key in ['standard_volume_rate', 'temperature_inlet', 'pressure_inlet', 'pressure_outlet']:
            if bc.get(key) is None:
                missing.append(f"boundary_condition.{key}")
        wm = self.parameters['working_medium']
        for key in ['k_index', 'R_index', 'rotation_speed']:
            if wm.get(key) is None:
                missing.append(f"working_medium.{key}")
        if not self.parameters['sections']:
            missing.append("sections")
        else:
            for i, section in enumerate(self.parameters['sections']):
                if not section.get('stages'):
                    missing.append(f"section_{i}.stages")
        if missing:
            return (False, f"参数不完整: {', '.join(missing)}")
        return (True, "参数完整有效")


# 2. 粒子群优化类
class PSOOptimizer:
    def __init__(self, config, objective_func):
        self.config = config
        self.objective_func = objective_func
        self.history = {'best_fitness': [], 'positions': []}
        self.n_particles = config.get('n_particles', 30)
        self.max_iter = config.get('max_iter', 100)
        self.w = config.get('inertia_weight', 0.7)
        self.c1 = config.get('cognitive_weight', 1.5)
        self.c2 = config.get('social_weight', 1.5)
        self.n_dim = config['sections'] - 1
        self.lb = np.ones(self.n_dim) * config['total_pressure_ratio'] * 0.85
        self.ub = np.ones(self.n_dim) * config['total_pressure_ratio'] * 1.20
        self.positions = np.random.uniform(self.lb, self.ub, (self.n_particles, self.n_dim))
        self.velocities = np.zeros((self.n_particles, self.n_dim))
        self.fitness = np.array([self.objective_func(self._get_full_solution(pos)) for pos in self.positions])
        self.pbest_positions = self.positions.copy()
        self.pbest_fitness = self.fitness.copy()
        self.gbest_position = self.pbest_positions[np.argmin(self.pbest_fitness)]
        self.gbest_fitness = np.min(self.pbest_fitness)

    def _get_full_solution(self, partial_solution):
        product = np.prod(partial_solution)
        last_ratio = self.config['total_pressure_ratio'] / product
        if last_ratio < 1.01:
            scale_factor = (self.config['total_pressure_ratio'] / 1.01) ** (1 / self.n_dim)
            partial_solution = partial_solution / np.prod(partial_solution) ** (1 / self.n_dim) * scale_factor
            last_ratio = 1.01
        elif last_ratio > self.config['total_pressure_ratio']:
            scale_factor = (self.config['total_pressure_ratio'] / self.config['total_pressure_ratio']) ** (
                    1 / self.n_dim)
            partial_solution = partial_solution / np.prod(partial_solution) ** (1 / self.n_dim) * scale_factor
            last_ratio = self.config['total_pressure_ratio']
        return np.append(partial_solution, last_ratio)

    def optimize(self):
        for iter in range(self.max_iter):
            r1 = np.random.rand(self.n_particles, self.n_dim)
            r2 = np.random.rand(self.n_particles, self.n_dim)
            self.velocities = (self.w * self.velocities +
                               self.c1 * r1 * (self.pbest_positions - self.positions) +
                               self.c2 * r2 * (self.gbest_position - self.positions))
            max_velocity = (self.ub - self.lb) * 0.2
            for d in range(self.n_dim):
                self.velocities[:, d] = np.clip(self.velocities[:, d], -max_velocity[d], max_velocity[d])
            self.positions += self.velocities
            self.positions = np.clip(self.positions, self.lb, self.ub)
            for i in range(self.n_particles):
                full_solution = self._get_full_solution(self.positions[i])
                new_fitness = self.objective_func(full_solution)
                if new_fitness < self.pbest_fitness[i]:
                    self.pbest_fitness[i] = new_fitness
                    self.pbest_positions[i] = self.positions[i].copy()
                    if new_fitness < self.gbest_fitness:
                        self.gbest_fitness = new_fitness
                        self.gbest_position = self.positions[i].copy()
            self.history['best_fitness'].append(self.gbest_fitness)
            self.history['positions'].append(self._get_full_solution(self.gbest_position).copy())
            if (iter + 1) % 10 == 0:
                st.write(f"Iteration {iter + 1}/{self.max_iter}, Best Fitness: {self.gbest_fitness:.4f}")
        return {
            'optimal_ratios': self._get_full_solution(self.gbest_position).tolist(),
            'optimal_fitness': float(self.gbest_fitness),
            'history': {'best_fitness': [float(f) for f in self.history['best_fitness']],
                        'positions': [pos.tolist() for pos in self.history['positions']]}
        }


# 3. 离心压缩机设计类
class CentrifugalCompressorDesign:
    def __init__(self, params):
        self.params = params
        self.q_v_sta = params["boundary_condition"]["standard_volume_rate"]
        self.p_in = params["boundary_condition"]["pressure_inlet"]
        self.t_in = params["boundary_condition"]["temperature_inlet"]
        self.p_out = params["boundary_condition"]["pressure_outlet"]
        self.t_w1 = 30
        self.mixture_params = {}
        self.segment_designs = []

    def calculate_mixture_properties(self):
        R_mix = self.params["working_medium"]["R_index"]
        k_mix = self.params["working_medium"]["k_index"]
        rho_mix = self.params["working_medium"]["rho_index"]
        q_v_in_act = self.q_v_sta * 0.101325 * (273 + self.t_in) / self.p_in / 273
        q_v_in = q_v_in_act * 1.03
        pressure_ratio = 1.03 * self.p_out / self.p_in - 0.03
        self.mixture_params = {
            'R_mix': R_mix, 'k_mix': k_mix, 'rho_mix': rho_mix,
            'q_v_in': q_v_in, 'pressure_ratio': pressure_ratio
        }

    def calculate_internal_power(self, pressure_ratios, segment_params):
        total_power = 0.0
        T_in = segment_params['T_d'][0]
        P_in = self.p_in
        m = self.mixture_params['q_v_in'] * self.mixture_params['rho_mix'] / 60
        R = self.mixture_params['R_mix']
        γ = self.mixture_params['k_mix']
        for i, ε in enumerate(pressure_ratios):
            η_poly = segment_params['eta_pol'][i]
            n = γ / (γ - η_poly * (γ - 1))
            W_poly = (n / (n - 1)) * R * T_in * (ε ** ((n - 1) / n) - 1)
            total_power += m * W_poly
            P_out = P_in * ε
            T_out = T_in * ε ** ((γ - 1) / γ)
            T_in = self.t_w1 + 273 + 12 if i < len(pressure_ratios) - 1 else T_out
            P_in = P_out
        return total_power / 1000

    def optimize_pressure_ratios(self, segment_params):
        pso_config = {
            'total_pressure_ratio': self.mixture_params['pressure_ratio'],
            'sections': segment_params['total_segments'],
            'n_particles': 20, 'max_iter': 50,
            'inertia_weight': 0.7, 'cognitive_weight': 1.5, 'social_weight': 1.5
        }

        def objective_func(pressure_ratios):
            return self.calculate_internal_power(pressure_ratios, segment_params)

        optimizer = PSOOptimizer(pso_config, objective_func)
        return optimizer.optimize()

    def segment_definition(self, cooling_times):
        Z_2 = cooling_times
        total_segments = Z_2 + 1
        eta_pol = [0.83, 0.81, 0.80, 0.78][:Z_2 + 1]
        if len(eta_pol) < Z_2 + 1:
            eta_pol.extend([eta_pol[-1]] * (Z_2 + 1 - len(eta_pol)))
        T_d = [self.t_in + 273] + [self.t_w1 + 273 + 12] * Z_2
        segment_params = {
            'cooling_times': cooling_times, 'Z_2': Z_2, 'total_segments': total_segments,
            'T_d': T_d, 'eta_pol': eta_pol, 'lambda_m': 0.9841
        }
        optimization_result = self.optimize_pressure_ratios(segment_params)
        pressure_ratio_seg = optimization_result['optimal_ratios']
        p_low_leng = [0.0039] * Z_2
        p_d = [self.p_in]
        for i in range(Z_2):
            p_d.append(p_d[i] * pressure_ratio_seg[i] - p_low_leng[i])
        segment_params.update({
            'pressure_ratio_seg': pressure_ratio_seg, 'p_d': p_d,
            'optimization_result': optimization_result
        })
        return segment_params

    def calculate_segment_performance(self, segment_params):
        q_v_d = [self.mixture_params['q_v_in']]
        for i in range(segment_params['Z_2']):
            q_v_d.append(q_v_d[0] * segment_params['p_d'][0] / segment_params['p_d'][i + 1] *
                         segment_params['T_d'][i + 1] / segment_params['T_d'][0])
        sigma = [self.mixture_params['k_mix'] / (self.mixture_params['k_mix'] - 1) * eta for eta in
                 segment_params['eta_pol']]
        W_pol = [s * self.mixture_params['R_mix'] * T * (pr ** (1 / s) - 1) for s, T, pr in
                 zip(sigma, segment_params['T_d'], segment_params['pressure_ratio_seg'])]
        segment_params.update({'q_v_d': q_v_d, 'sigma': sigma, 'W_pol': W_pol})
        return segment_params

    def impeller_design(self, segment_params):
        beta_2A_d = [math.pi / 2, math.pi * 5 / 18, math.pi * 5 / 18, math.pi * 5 / 18]
        phi_2r_d = [0.33, 0.27, 0.26, 0.25]
        Z_z = [24, 24, 24, 24]
        n_segments = segment_params['total_segments']
        beta_2A_d = beta_2A_d[:n_segments] + [beta_2A_d[-1]] * max(0, n_segments - len(beta_2A_d))
        phi_2r_d = phi_2r_d[:n_segments] + [phi_2r_d[-1]] * max(0, n_segments - len(phi_2r_d))
        Z_z = Z_z[:n_segments] + [Z_z[-1]] * max(0, n_segments - len(Z_z))
        phi_2u_d = [1 - math.pi / z * math.sin(beta) - phi_2r / math.tan(beta) for z, beta, phi_2r in
                    zip(Z_z, beta_2A_d, phi_2r_d)]
        loss_d = [0.01456, 0.03103, 0.03587, 0.04809][:n_segments] + [0.04809] * max(0, n_segments - 4)
        eta_h = [eta_pol * (1 + loss) for eta_pol, loss in zip(segment_params['eta_pol'], loss_d)]
        psi_d = [eta * phi for eta, phi in zip(eta_h, phi_2u_d)]
        u_2_d = [285, 285, 250, 250][:n_segments] + [250] * max(0, n_segments - 4)
        i_d = [max(1, math.ceil(W_pol / psi / u ** 2 - 0.1)) for W_pol, psi, u in
               zip(segment_params['W_pol'], psi_d, u_2_d)]
        u_2_d_adj = [math.sqrt(W_pol / i / psi) for W_pol, i, psi in zip(segment_params['W_pol'], i_d, psi_d)]
        return {
            'beta_2A_d': beta_2A_d, 'phi_2r_d': phi_2r_d, 'Z_z': Z_z,
            'phi_2u_d': phi_2u_d, 'eta_h': eta_h, 'psi_d': psi_d,
            'u_2_d_initial': u_2_d, 'i_d': i_d, 'u_2_d_adjusted': u_2_d_adj
        }

    def rotor_dynamics_analysis(self, segment_params, impeller_params):
        b_div_D = 0.06536
        tau_2_d = [0.9236, 0.90027, 0.88919, 0.88603]
        n_segments = segment_params['total_segments']
        tau_2_d = tau_2_d[:n_segments] + [tau_2_d[-1]] * max(0, n_segments - 4)
        c_2r = [u * phi for u, phi in zip(impeller_params['u_2_d_adjusted'], impeller_params['phi_2r_d'])]
        alpha_2 = [math.atan(phi_r / phi_u) for phi_r, phi_u in
                   zip(impeller_params['phi_2r_d'], impeller_params['phi_2u_d'])]
        c_2 = [c_r / math.sin(alpha) for c_r, alpha in zip(c_2r, alpha_2)]
        delta_T2_j = [(self.mixture_params['k_mix'] - 1) / self.mixture_params['k_mix'] / self.mixture_params['R_mix'] *
                      (W_pol / eta_pol - c ** 2 / 2) for W_pol, eta_pol, c in
                      zip(segment_params['W_pol'], segment_params['eta_pol'], c_2)]
        K_v2_j = [(1 + delta_T / T) ** (sigma - 1) for delta_T, T, sigma in
                  zip(delta_T2_j, segment_params['T_d'], segment_params['sigma'])]
        n = 33.9 * math.sqrt(K_v2_j[0] * b_div_D * tau_2_d[0] *
                             impeller_params['phi_2r_d'][0] * impeller_params['u_2_d_adjusted'][0] ** 3 /
                             self.mixture_params['q_v_in'] * 60)
        b_div_D_j = [q_v / 60 / K_v / tau / u ** 3 / phi_r * (n / 33.9) ** 2 for q_v, K_v, tau, u, phi_r in zip(
            segment_params['q_v_d'], K_v2_j, tau_2_d, impeller_params['u_2_d_adjusted'], impeller_params['phi_2r_d'])]
        D_2_j = [60 * u / (math.pi * n) for u in impeller_params['u_2_d_adjusted']]
        delta_calc = [4e-3] * n_segments
        tau_2_j_real = [1 - Z_z * delta / (math.pi * D * math.sin(beta)) for Z_z, delta, D, beta in zip(
            impeller_params['Z_z'], delta_calc, D_2_j, impeller_params['beta_2A_d'])]
        P_j = [self.mixture_params['q_v_in'] * self.mixture_params['rho_mix'] * W_pol / 60 for W_pol in
               segment_params['W_pol']]
        return {
            'n': n, 'b_div_D_j': b_div_D_j, 'D_2_j': D_2_j,
            'tau_2_j_real': tau_2_j_real, 'P_j': P_j, 'total_power': sum(P_j)
        }

    def initialize_stage(self):
        try:
            result = subprocess.run(
                ['python', 'D:\\programming\\forVSCode\\NEW\\prog\\InitializeStage.py'],
                capture_output=True, text=True, timeout=60
            )
            matches = re.findall(r'(?:建议选取中冷|或中冷)\s*(\d+)\s*次', result.stdout)
            return [int(m) for m in matches] if matches else [1, 2]
        except Exception as e:
            st.warning(f"初始化阶段失败: {e}")
            return [1, 2]

    def save_parameters(self, file_path):
        try:
            design_results = {
                "initial_parameters": {
                    "standard_volume_rate": self.q_v_sta,
                    "inlet_pressure": self.p_in,
                    "inlet_temperature": self.t_in,
                    "outlet_pressure": self.p_out,
                    "cooling_water_temperature": self.t_w1
                },
                "mixture_properties": self.mixture_params,
                "segment_designs": self.segment_designs
            }
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(design_results, f, indent=4, ensure_ascii=False)
            return (True, f"设计结果已保存到 {file_path}")
        except Exception as e:
            return (False, f"保存失败: {str(e)}")

    def select_strategy(self):
        if not self.segment_designs:
            return None
        min_power = min(d['stage_params']['total_power'] for d in self.segment_designs)
        for i, d in enumerate(self.segment_designs):
            if d['stage_params']['total_power'] == min_power:
                return i
        return None

    def run_design(self, progress_callback=None):
        self.calculate_mixture_properties()
        if progress_callback:
            progress_callback(1)
        cooling_options = self.initialize_stage()
        if not cooling_options:
            cooling_options = [1]
        if progress_callback:
            progress_callback(1)
        if progress_callback:
            progress_callback(2)
        for cooling_times in cooling_options:
            try:
                segment_params = self.segment_definition(cooling_times)
                if progress_callback:
                    progress_callback(2)
                segment_params = self.calculate_segment_performance(segment_params)
                impeller_params = self.impeller_design(segment_params)
                stage_params = self.rotor_dynamics_analysis(segment_params, impeller_params)
                self.segment_designs.append({
                    'cooling_times': cooling_times,
                    'segment_params': segment_params,
                    'impeller_params': impeller_params,
                    'stage_params': stage_params
                })
            except Exception as e:
                st.warning(f"生成冷却次数为{cooling_times}的方案失败: {e}")
                continue
        if progress_callback:
            progress_callback(3)
        best_index = self.select_strategy()
        if best_index is not None and 0 <= best_index < len(self.segment_designs):
            self.segment_designs[best_index]['is_best'] = True
        else:
            st.warning("警告：未找到有效方案，无法标记最佳方案")


# 4. 报告生成函数（完全保留）
def generate_design_report(json_file_path, output_docx_path):
    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    doc = Document()
    title = doc.add_heading('离心压缩机热力设计报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    time_para = doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
    time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 初始参数表格
    doc.add_heading('一、初始设计参数', level=1)
    initial_params = data['initial_parameters']
    rows_data = [
        ('标准体积流量', f"{initial_params['standard_volume_rate']} m³/min"),
        ('进口压力', f"{initial_params['inlet_pressure']} MPa"),
        ('进口温度', f"{initial_params['inlet_temperature']} K"),
        ('出口压力', f"{initial_params['outlet_pressure']} MPa"),
        ('冷却水温度', f"{initial_params['cooling_water_temperature']} °C")
    ]
    initial_table = doc.add_table(rows=len(rows_data) + 1, cols=2)
    initial_table.style = 'Light Shading'
    initial_table.cell(0, 0).text = '参数名称'
    initial_table.cell(0, 1).text = '参数值'
    for i, (param_name, param_value) in enumerate(rows_data, start=1):
        if i < len(initial_table.rows):
            initial_table.cell(i, 0).text = param_name
            initial_table.cell(i, 1).text = param_value

    # 物性参数表格
    doc.add_heading('二、混合气体物性参数', level=1)
    mixture_props = data['mixture_properties']
    mixture_rows = [
        ('气体常数', f"{mixture_props['R_mix']:.2f} J/(kg·K)"),
        ('等熵指数', f"{mixture_props['k_mix']:.4f}"),
        ('密度', f"{mixture_props['rho_mix']:.4f} kg/m³"),
        ('进气体积流量', f"{mixture_props['q_v_in']:.2f} m³/min"),
        ('总压比', f"{mixture_props['pressure_ratio']:.4f}")
    ]
    mixture_table = doc.add_table(rows=len(mixture_rows) + 1, cols=2)
    mixture_table.style = 'Light Shading'
    mixture_table.cell(0, 0).text = '物性参数'
    mixture_table.cell(0, 1).text = '数值'
    for i, (param_name, param_value) in enumerate(mixture_rows, start=1):
        if i < len(mixture_table.rows):
            mixture_table.cell(i, 0).text = param_name
            mixture_table.cell(i, 1).text = param_value

    # 方案对比表格
    doc.add_heading('三、分段设计方案比较', level=1)
    summary_table = doc.add_table(rows=len(data['segment_designs']) + 1, cols=5)
    summary_table.style = 'Light Shading'
    summary_headers = ['冷却次数', '总内功率 (kW)', '转速 (r/s)', '叶轮直径 (m)', '是否最佳方案']
    for col, header in enumerate(summary_headers):
        summary_table.cell(0, col).text = header
    for i, design in enumerate(data['segment_designs'], start=1):
        cooling_times = design['cooling_times']
        total_power = design['stage_params']['total_power'] / 1000
        speed = design['stage_params']['n']
        impeller_diameter = design['stage_params']['D_2_j'][0]
        is_best = "是" if design.get('is_best', False) else "否"
        summary_table.cell(i, 0).text = str(cooling_times)
        summary_table.cell(i, 1).text = f"{total_power:.2f}"
        summary_table.cell(i, 2).text = f"{speed:.2f}"
        summary_table.cell(i, 3).text = f"{impeller_diameter:.4f}"
        summary_table.cell(i, 4).text = is_best

    # 结论
    doc.add_heading('四、设计结论与建议', level=1)
    best_design = next((d for d in data['segment_designs'] if d.get('is_best')), None)
    if best_design:
        conclusion = doc.add_paragraph()
        conclusion.add_run("设计结论: ").bold = True
        conclusion.add_run(
            f"推荐中冷{best_design['cooling_times']}次方案，总功率{best_design['stage_params']['total_power'] / 1000:.2f}kW")
    doc.save(output_docx_path)
    return True


# ===================== 3. 新增API接口逻辑（不影响原有前端） =====================
def handle_api_request():
    """处理API请求，返回JSON格式结果"""
    query_params = st.query_params.to_dict()

    # 检测是否为API调用（url中包含?api=1）
    if "api" in query_params and query_params["api"][0] == "1":
        try:
            # 从URL参数获取设计参数
            standard_volume_rate = float(query_params.get("standard_volume_rate", [500])[0])
            p_in = float(query_params.get("p_in", [0.1])[0])
            t_in = float(query_params.get("t_in", [300])[0])
            p_out = float(query_params.get("p_out", [0.6])[0])
            medium = query_params.get("medium", ["氮气"])[0]
            section_type = query_params.get("section_type", ["high_efficiency"])[0]
            num_stages = int(query_params.get("num_stages", [3])[0])

            # 构造参数字符串（适配原解析函数）
            design_params = f"""
            标准体积流量={standard_volume_rate}m³/h;
            进口压力={p_in}MPa;
            进口温度={t_in}K;
            出口压力={p_out}MPa;
            组分1=CoolProp代名为={WORKING_MEDIA_OPTIONS[medium]};
            组分1摩尔分数=100%;
            """

            # 执行核心计算逻辑
            generator = ParameterGenerator()
            parse_ok, parse_msg = generator.parse_boundary_conditions(design_params)
            if not parse_ok:
                st.write(json.dumps({"code": 500, "error": parse_msg}))
                st.stop()

            medium_ok, medium_msg = generator.calculate_working_medium_properties(design_params)
            if not medium_ok:
                st.write(json.dumps({"code": 500, "error": medium_msg}))
                st.stop()

            generator.generate_section_stage(section_type, num_stages=num_stages)
            validate_ok, validate_msg = generator.validate_parameters()
            if not validate_ok:
                st.write(json.dumps({"code": 500, "error": validate_msg}))
                st.stop()

            # 压缩机设计
            compressor = CentrifugalCompressorDesign(generator.parameters)
            compressor.run_design()

            # 整理结果为JSON
            result_list = []
            for d in compressor.segment_designs:
                result_list.append({
                    "cooling_times": d["cooling_times"],
                    "total_segments": d["segment_params"]["total_segments"],
                    "total_power_kW": round(d["stage_params"]["total_power"] / 1000, 2),
                    "speed_rps": round(d["stage_params"]["n"], 2),
                    "impeller_diameter_m": round(d["stage_params"]["D_2_j"][0], 4),
                    "is_best": d.get("is_best", False)
                })

            # 生成报告（可选）
            result_path = os.path.join(STORAGE["design_results"],
                                       f"api_result_{datetime.now().strftime('%Y%m%d%H%M%S')}.json")
            compressor.save_parameters(result_path)
            report_path = os.path.join(STORAGE["reports"], f"api_report_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx")
            generate_design_report(result_path, report_path)

            # 返回API结果
            api_result = {
                "code": 200,
                "message": "设计成功",
                "data": {
                    "input_params": {
                        "standard_volume_rate": standard_volume_rate,
                        "p_in": p_in,
                        "t_in": t_in,
                        "p_out": p_out,
                        "medium": medium
                    },
                    "design_results": result_list,
                    "report_path": report_path
                }
            }
            st.write(json.dumps(api_result, ensure_ascii=False))
            st.stop()  # 停止渲染前端界面

        except Exception as e:
            st.write(json.dumps({"code": 500, "error": str(e)}))
            st.stop()


# ===================== 4. Streamlit前端（替换为参数选择框，删除AI） =====================
def main():
    # 先处理API请求
    handle_api_request()

    st.set_page_config(page_title="离心压缩机设计系统", page_icon="🔧", layout="wide")

    # 会话状态初始化
    if "design_results" not in st.session_state:
        st.session_state.design_results = None
    if "report_path" not in st.session_state:
        st.session_state.report_path = None
    if "stage" not in st.session_state:
        st.session_state.stage = 0

    # 侧边栏参数选择框（核心修改：替换文字输入）
    with st.sidebar:
        st.title("🔧 设计参数配置")
        st.markdown("### 边界条件")
        standard_volume_rate = st.number_input("标准体积流量 (m³/h)", min_value=100.0, max_value=10000.0, value=500.0,
                                               step=10.0)
        p_in = st.number_input("进口压力 (MPa)", min_value=0.01, max_value=10.0, value=0.1, step=0.01)
        t_in = st.number_input("进口温度 (K)", min_value=200.0, max_value=500.0, value=300.0, step=1.0)
        p_out = st.number_input("出口压力 (MPa)", min_value=0.02, max_value=20.0, value=0.6, step=0.01)

        st.markdown("### 工质配置")
        medium = st.selectbox("工质类型", list(WORKING_MEDIA_OPTIONS.keys()), index=1)  # 默认氮气
        mole_fraction = st.slider("摩尔分数 (%)", min_value=90.0, max_value=100.0, value=100.0, step=0.1)

        st.markdown("### 段/级配置")
        section_type = st.selectbox("段类型", list(SECTION_TEMPLATES.keys()), index=1)  # 默认高效型
        num_stages = st.slider("每段级数", min_value=1, max_value=5, value=3)

    # 主页面
    st.title("离心压缩机热力设计系统")
    st.markdown("### API调用示例")
    st.code(f"""
    https://你的网址.streamlit.app?api=1&standard_volume_rate={standard_volume_rate}&p_in={p_in}&t_in={t_in}&p_out={p_out}&medium={medium}
    """)

    # 提交设计
    if st.button("📤 提交设计", type="primary", use_container_width=True):
        progress_bar = st.progress(0, text=STAGE_MAP[0])
        st.session_state.stage = 0

        try:
            # 构造参数字符串（替换原AI提取逻辑）
            progress_bar.progress(20, text=STAGE_MAP[0])
            design_params = f"""
            标准体积流量={standard_volume_rate}m³/h;
            进口压力={p_in}MPa;
            进口温度={t_in}K;
            出口压力={p_out}MPa;
            组分1=CoolProp代名为={WORKING_MEDIA_OPTIONS[medium]};
            组分1摩尔分数={mole_fraction}%;
            """
            st.success(f"✅ 参数配置完成：{medium} | 进口压力 {p_in}MPa | 进口温度 {t_in}K")

            # 处理数据（原逻辑）
            progress_bar.progress(40, text=STAGE_MAP[1])
            generator = ParameterGenerator()
            parse_ok, parse_msg = generator.parse_boundary_conditions(design_params)
            if not parse_ok:
                st.error(f"❌ {parse_msg}")
                return

            medium_ok, medium_msg = generator.calculate_working_medium_properties(design_params)
            if not medium_ok:
                st.error(f"❌ {medium_msg}")
                return

            generator.generate_section_stage(section_type, num_stages=num_stages)
            validate_ok, validate_msg = generator.validate_parameters()
            if not validate_ok:
                st.error(f"❌ {validate_msg}")
                return
            st.success(f"✅ 数据处理完成：{parse_msg} | {medium_msg}")

            # 优化设计（原逻辑）
            progress_bar.progress(70, text=STAGE_MAP[2])
            compressor = CentrifugalCompressorDesign(generator.parameters)

            def update_progress(stage_num):
                st.session_state.stage = stage_num
                progress_bar.progress(20 + (stage_num - 1) * 20, text=STAGE_MAP[stage_num])

            compressor.run_design(progress_callback=update_progress)
            if not compressor.segment_designs:
                st.error("❌ 未生成设计方案")
                return
            st.session_state.design_results = compressor.segment_designs

            # 生成报告（原逻辑，删除AI评估）
            progress_bar.progress(90, text=STAGE_MAP[3])
            result_path = os.path.join(STORAGE["design_results"],
                                       f"result_{datetime.now().strftime('%Y%m%d%H%M%S')}.json")
            compressor.save_parameters(result_path)
            report_path = os.path.join(STORAGE["reports"], f"report_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx")
            generate_design_report(result_path, report_path)
            st.session_state.report_path = report_path

            # 完成
            progress_bar.progress(100, text=STAGE_MAP[4])
            st.success("🎉 设计完成！")

        except Exception as e:
            st.error(f"❌ 设计异常：{str(e)}")
            progress_bar.progress(100, text=STAGE_MAP[5])

    # 结果展示
    st.divider()
    st.subheader("设计结果展示")
    if st.session_state.design_results:
        # 方案对比
        st.markdown("### 1. 设计方案对比")
        df = pd.DataFrame([{
            "冷却次数": d["cooling_times"],
            "总段数": d["segment_params"]["total_segments"],
            "总功率(kW)": round(d["stage_params"]["total_power"] / 1000, 2),
            "转速(r/s)": round(d["stage_params"]["n"], 2),
            "叶轮直径(m)": round(d["stage_params"]["D_2_j"][0], 4),
            "最优方案": "✅" if d.get("is_best") else "❌"
        } for d in st.session_state.design_results])
        st.dataframe(df, use_container_width=True)

        # 报告下载
        st.markdown("### 2. 报告下载")
        if st.session_state.report_path and os.path.exists(st.session_state.report_path):
            with open(st.session_state.report_path, 'rb') as f:
                st.download_button(
                    "📄 下载DOCX报告",
                    f,
                    file_name=os.path.basename(st.session_state.report_path),
                    use_container_width=True
                )


# ===================== 4. 主函数执行 =====================
if __name__ == "__main__":
    main()